"""
generate_persona_decision_report.py
------------------------------------
Generates a detailed persona decision report across Journeys A, B, C.
For each journey, picks one reorderer and one lapser.
Produces:
  - reports/persona_decisions/Persona_Decision_Report.docx
  - reports/persona_decisions/persona_decision_data.json  (raw machine-readable)
"""

from __future__ import annotations

import json
import os
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Paths ────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent.parent
DATA = ROOT / "data" / "population"
OUT_DIR = ROOT / "reports" / "persona_decisions"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_OUT = OUT_DIR / "Persona_Decision_Report.docx"
JSON_OUT  = OUT_DIR / "persona_decision_data.json"

# ── Persona pairs to analyse ─────────────────────────────────────────────────
PAIRS = {
    "A": {
        "journey_file": "journey_A_results.json",
        "journey_name": "Journey A — Nutrimix Repeat Purchase (Age 2–6)",
        "reorderer": "Nasreen-Pune-Mom-34",
        "lapser":    "Ayesha-Delhi-Mom-37",
        "decision_ticks": [20, 60],
        "product": "littlejoys",
    },
    "B": {
        "journey_file": "journey_B_results.json",
        "journey_name": "Journey B — Magnesium Gummies Growth (Age 3–8)",
        "reorderer": "Om-Nagpur-Dad-34",
        "lapser":    "Devi-Dehradun-Mom-32",
        "decision_ticks": [35, 60],
        "product": "littlejoys",
    },
    "C": {
        "journey_file": "journey_C_results.json",
        "journey_name": "Journey C — Nutrimix 7–14 Expansion",
        "reorderer": "Ayesha-Delhi-Mom-37",
        "lapser":    "Om-Nagpur-Dad-34",
        "decision_ticks": [28, 60],
        "product": "littlejoys",
    },
}

# ── Colour palette ────────────────────────────────────────────────────────────
CLR = {
    "void":      RGBColor(0x05, 0x05, 0x05),
    "heading":   RGBColor(0x1A, 0x1A, 0x1A),
    "body":      RGBColor(0x2C, 0x2C, 0x2C),
    "dim":       RGBColor(0x5E, 0x5E, 0x5E),
    "signal":    RGBColor(0x1A, 0x7A, 0x1A),   # muted green for print
    "reorder":   RGBColor(0x0D, 0x6E, 0x0D),
    "lapse":     RGBColor(0xC0, 0x39, 0x2B),
    "neutral":   RGBColor(0x27, 0x6F, 0xB5),
    "row_even":  RGBColor(0xF7, 0xF7, 0xF7),
    "row_odd":   RGBColor(0xFF, 0xFF, 0xFF),
    "header_bg": RGBColor(0x1A, 0x1A, 0x1A),
    "header_fg": RGBColor(0xFF, 0xFF, 0xFF),
}

# ── Data helpers ──────────────────────────────────────────────────────────────

def load_json(path: Path) -> dict | list:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def build_cohort_lookup(cohort_data) -> dict:
    """Build name→persona lookup from simulatte_cohort_final.json."""
    lookup = {}
    personas = cohort_data.get("personas", [])
    for p in personas:
        demo = p.get("demographic_anchor", {})
        name = demo.get("name", "")
        city = demo.get("location", {}).get("city", "")
        key = f"{name}-{city}"
        lookup[key] = p
    return lookup


def get_persona_profile(cohort_lookup: dict, persona_id: str) -> dict | None:
    """Match journey persona_id like 'Nasreen-Pune-Mom-34' to cohort profile."""
    parts = persona_id.split("-")
    if len(parts) >= 2:
        key = f"{parts[0]}-{parts[1]}"
        return cohort_lookup.get(key)
    return None


def get_log(logs: list, persona_id: str) -> dict | None:
    return next((l for l in logs if l.get("persona_id") == persona_id), None)


def get_snap(log: dict, tick: int) -> dict | None:
    return next((s for s in log.get("snapshots", []) if s["tick"] == tick), None)


def get_primary_trust(brand_trust: dict) -> tuple[str, float]:
    if not brand_trust:
        return ("unknown", 0.0)
    return list(brand_trust.items())[0]


def get_competitor_trusts(brand_trust: dict) -> dict:
    items = list(brand_trust.items())
    return {k: v for k, v in items[1:]} if len(items) > 1 else {}


def extract_stimuli_window(log: dict, after_tick: int) -> list[dict]:
    """Return all stimuli fired after after_tick."""
    results = []
    for snap in log.get("snapshots", []):
        t = snap["tick"]
        if t <= after_tick:
            continue
        pr = snap.get("perception_results")
        if pr and isinstance(pr, list) and pr:
            p = pr[0]
            results.append({
                "tick": t,
                "stimulus_id": p.get("stimulus_id"),
                "emotional_valence": p.get("emotional_valence"),
                "importance": p.get("importance"),
                "reflection_trigger": p.get("reflection_trigger", False),
            })
    return results


def build_persona_json(
    persona_id: str,
    log: dict,
    profile: dict | None,
    decision_ticks: list[int],
    product_key: str,
) -> dict:
    """Build the full machine-readable record for one persona."""
    snaps = {s["tick"]: s for s in log.get("snapshots", [])}

    # Static profile
    profile_data = {}
    if profile:
        demo = profile.get("demographic_anchor", {})
        attrs = profile.get("attributes", {})
        derived = profile.get("derived_insights", {})
        beh = profile.get("behavioural_tendencies", {})

        def flatten_attrs(group: dict) -> dict:
            return {k: round(v["value"], 4) for k, v in group.items() if "value" in v}

        profile_data = {
            "name": demo.get("name"),
            "age": demo.get("age"),
            "gender": demo.get("gender"),
            "city": demo.get("location", {}).get("city"),
            "urban_tier": demo.get("location", {}).get("urban_tier"),
            "household_structure": demo.get("household", {}).get("structure"),
            "household_size": demo.get("household", {}).get("size"),
            "income_bracket": demo.get("household", {}).get("income_bracket"),
            "education": demo.get("education"),
            "employment": demo.get("employment"),
            "attributes": {
                "values":     flatten_attrs(attrs.get("values", {})),
                "social":     flatten_attrs(attrs.get("social", {})),
                "lifestyle":  flatten_attrs(attrs.get("lifestyle", {})),
                "psychology": flatten_attrs(attrs.get("psychology", {})),
            },
            "derived_insights": {
                k: v for k, v in derived.items()
                if not isinstance(v, (dict, list))
            },
            "behavioural_tendencies": {
                k: v for k, v in beh.items()
                if not isinstance(v, (dict, list))
            },
        }

    # Decision snapshots
    decision_data = {}
    for tick in decision_ticks:
        snap = snaps.get(tick)
        if not snap:
            continue
        dr = snap.get("decision_result") or {}
        bt = snap.get("brand_trust", {})
        pk, pt = get_primary_trust(bt)
        decision_data[f"tick_{tick}"] = {
            "tick": tick,
            "brand_trust": {
                "primary": {"brand": pk, "value": round(pt, 4)},
                "competitors": {k: round(v, 4) for k, v in get_competitor_trusts(bt).items()},
            },
            "memories_count": snap.get("memories_count", 0),
            "cumulative_salience": round(snap.get("cumulative_salience", 0), 4),
            "reflected": snap.get("reflected", False),
            "decision": {
                "outcome": dr.get("decision"),
                "confidence": dr.get("confidence"),
                "willingness_to_pay_inr": dr.get("willingness_to_pay_inr"),
                "implied_purchase": dr.get("implied_purchase", False),
                "key_drivers": dr.get("key_drivers", []),
                "objections": dr.get("objections", []),
                "follow_up_action": dr.get("follow_up_action"),
                "reasoning_trace": dr.get("reasoning_trace", []),
            },
        }

    # Compute deltas between consecutive decision ticks
    deltas = {}
    for i in range(1, len(decision_ticks)):
        t_prev = decision_ticks[i - 1]
        t_curr = decision_ticks[i]
        key = f"tick_{t_prev}_to_{t_curr}"
        d_prev = decision_data.get(f"tick_{t_prev}", {})
        d_curr = decision_data.get(f"tick_{t_curr}", {})
        if d_prev and d_curr:
            trust_prev = d_prev["brand_trust"]["primary"]["value"]
            trust_curr = d_curr["brand_trust"]["primary"]["value"]
            mem_prev   = d_prev["memories_count"]
            mem_curr   = d_curr["memories_count"]
            deltas[key] = {
                "trust_delta": round(trust_curr - trust_prev, 4),
                "trust_prev": trust_prev,
                "trust_curr": trust_curr,
                "memories_added": mem_curr - mem_prev,
                "decision_changed": d_prev["decision"]["outcome"] != d_curr["decision"]["outcome"],
                "competitor_brands_gained": list(
                    set(d_curr["brand_trust"]["competitors"].keys()) -
                    set(d_prev["brand_trust"]["competitors"].keys())
                ),
                "competitor_trust_deltas": {
                    k: round(d_curr["brand_trust"]["competitors"].get(k, 0) - d_prev["brand_trust"]["competitors"].get(k, 0), 4)
                    for k in set(list(d_prev["brand_trust"]["competitors"].keys()) + list(d_curr["brand_trust"]["competitors"].keys()))
                },
            }

    # Post-first-decision stimulus log
    post_stimuli = extract_stimuli_window(log, decision_ticks[0])

    return {
        "persona_id": persona_id,
        "display_name": log.get("display_name"),
        "reordered": log.get("reordered"),
        "final_decision_outcome": (log.get("final_decision") or {}).get("decision"),
        "profile": profile_data,
        "decision_snapshots": decision_data,
        "deltas": deltas,
        "post_first_decision_stimuli": post_stimuli,
    }


# ── DOCX helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = CLR["heading"]
    return p


def add_para(doc: Document, text: str, bold=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or CLR["body"]
    return p


def add_json_block(doc: Document, data: dict, title: str = ""):
    """Add a JSON code block as a styled paragraph."""
    if title:
        add_para(doc, title, bold=True, size=9, color=CLR["dim"])
    json_str = json.dumps(data, indent=2, ensure_ascii=False)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(json_str)
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)
    run.font.color.rgb = CLR["body"]
    return p


def add_kv_table(doc: Document, rows: list[tuple], title: str = ""):
    """Add a two-column key-value table."""
    if title:
        add_para(doc, title, bold=True, size=9, color=CLR["dim"])
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i]
        k_cell = row.cells[0]
        v_cell = row.cells[1]
        bg = "F7F7F7" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(k_cell, bg)
        set_cell_bg(v_cell, bg)
        k_run = k_cell.paragraphs[0].add_run(str(k))
        k_run.font.size = Pt(8.5)
        k_run.bold = True
        k_run.font.color.rgb = CLR["dim"]
        v_run = v_cell.paragraphs[0].add_run(str(v))
        v_run.font.size = Pt(8.5)
        v_run.font.color.rgb = CLR["body"]
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(11)
    return tbl


def add_delta_callout(doc: Document, deltas: dict):
    """Highlight the key deltas between decision ticks in a readable table."""
    for window, d in deltas.items():
        ticks = window.replace("tick_", "").replace("_to_", " → tick ")
        add_para(doc, f"Δ  {ticks}", bold=True, size=9, color=CLR["dim"])
        rows = [
            ("Trust delta",         f"{d['trust_prev']:.4f}  →  {d['trust_curr']:.4f}  ({d['trust_delta']:+.4f})"),
            ("Memories added",      str(d["memories_added"])),
            ("Decision changed",    "YES" if d["decision_changed"] else "no"),
            ("New competitors",     ", ".join(d["competitor_brands_gained"]) or "none"),
            ("Competitor Δ trust",  json.dumps(d["competitor_trust_deltas"]) if d["competitor_trust_deltas"] else "none"),
        ]
        add_kv_table(doc, rows)
        doc.add_paragraph()


def add_stimulus_table(doc: Document, stimuli: list[dict]):
    """Table of post-decision stimuli with valence direction."""
    if not stimuli:
        add_para(doc, "No stimuli fired in post-decision window.", color=CLR["dim"], size=8.5)
        return
    tbl = doc.add_table(rows=1 + len(stimuli), cols=5)
    tbl.style = "Table Grid"
    headers = ["Tick", "Stimulus ID", "Valence", "Importance", "Direction"]
    header_row = tbl.rows[0]
    for j, h in enumerate(headers):
        set_cell_bg(header_row.cells[j], "1A1A1A")
        run = header_row.cells[j].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, stim in enumerate(stimuli):
        row = tbl.rows[i + 1]
        valence = stim.get("emotional_valence", 0)
        direction = "⬆ boost" if valence > 0.4 else ("⬇ erosion" if valence < 0 else "neutral")
        bg = "FFF5F5" if valence < 0 else ("F5FFF5" if valence > 0.4 else "FFFFFF")
        vals = [
            str(stim["tick"]),
            stim.get("stimulus_id", ""),
            f"{valence:+.2f}",
            f"{stim.get('importance', 0):.2f}",
            direction,
        ]
        for j, val in enumerate(vals):
            set_cell_bg(row.cells[j], bg)
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if valence < 0:
                run.font.color.rgb = CLR["lapse"]
            elif valence > 0.4:
                run.font.color.rgb = CLR["reorder"]
            else:
                run.font.color.rgb = CLR["body"]


def add_decision_section(doc: Document, snap_data: dict, tick: int, label: str):
    """Full detail section for one decision tick."""
    dec = snap_data.get("decision", {})
    bt  = snap_data.get("brand_trust", {})
    pk  = bt.get("primary", {}).get("brand", "?")
    pt  = bt.get("primary", {}).get("value", 0)
    comps = bt.get("competitors", {})

    add_para(doc, f"Decision at tick {tick} — {label}", bold=True, size=11,
             color=CLR["reorder"] if "reorder" in label.lower() else CLR["lapse"])
    doc.add_paragraph()

    # Core state table
    comp_str = ", ".join(f"{k}: {v:.4f}" for k, v in comps.items()) if comps else "none"
    rows = [
        ("Tick",                    tick),
        (f"{pk} brand trust",       f"{pt:.4f}"),
        ("Competitor trusts",        comp_str),
        ("Memories accumulated",    snap_data.get("memories_count", 0)),
        ("Cumulative salience",     snap_data.get("cumulative_salience", 0)),
        ("Has reflected",           snap_data.get("reflected", False)),
        ("Decision outcome",        dec.get("outcome", "—")),
        ("Confidence",              dec.get("confidence", "—")),
        ("Willingness to pay (Rs)", dec.get("willingness_to_pay_inr", "—")),
        ("Implied purchase",        dec.get("implied_purchase", False)),
    ]
    add_kv_table(doc, rows)
    doc.add_paragraph()

    # Key drivers
    drivers = dec.get("key_drivers", [])
    if drivers:
        add_para(doc, "Key drivers:", bold=True, size=9, color=CLR["dim"])
        for d in drivers:
            add_para(doc, f"+ {d}", size=9, indent=0.5, color=CLR["reorder"])

    # Objections
    objections = dec.get("objections", [])
    if objections:
        add_para(doc, "Objections:", bold=True, size=9, color=CLR["dim"])
        for o in objections:
            add_para(doc, f"− {o}", size=9, indent=0.5, color=CLR["lapse"])

    doc.add_paragraph()

    # Reasoning trace
    add_para(doc, "Reasoning trace (5 steps):", bold=True, size=9, color=CLR["dim"])
    for step in dec.get("reasoning_trace", []):
        add_para(doc, step, size=9, indent=0.5, color=CLR["body"])
        doc.add_paragraph()

    # Follow-up action
    if dec.get("follow_up_action"):
        add_para(doc, f"Follow-up action: {dec['follow_up_action']}", size=8.5, indent=0.5, color=CLR["dim"])
    doc.add_paragraph()


def add_full_json_dump(doc: Document, persona_rec: dict):
    """Add the complete machine-readable JSON for a persona."""
    add_para(doc, "Full data record (JSON)", bold=True, size=9, color=CLR["dim"])
    doc.add_paragraph()
    # Split into logical sections so the document stays readable
    sections = {
        "profile": persona_rec.get("profile", {}),
        "decision_snapshots": persona_rec.get("decision_snapshots", {}),
        "deltas": persona_rec.get("deltas", {}),
        "post_first_decision_stimuli": persona_rec.get("post_first_decision_stimuli", []),
    }
    for section_name, section_data in sections.items():
        add_para(doc, f"── {section_name}", bold=True, size=8.5, color=CLR["neutral"])
        json_str = json.dumps(section_data, indent=2, ensure_ascii=False)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(json_str)
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        doc.add_paragraph()


# ── Main build ────────────────────────────────────────────────────────────────

def build():
    cohort = load_json(DATA / "simulatte_cohort_final.json")
    cohort_lookup = build_cohort_lookup(cohort)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Title page
    doc.add_heading("LittleJoys — Persona Decision Deep-Dive", 0)
    add_para(doc, "Before vs After: Full data record at every decision tick across 6 personas across Journeys A, B and C.", size=11)
    add_para(doc, "Generated by Simulatte · Decision Infrastructure", size=9, color=CLR["dim"])
    doc.add_page_break()

    # Master JSON record for the raw output file
    master_json: dict = {}

    # ── Process each journey ──────────────────────────────────────────────────
    for journey_key, cfg in PAIRS.items():
        journey_data = load_json(DATA / cfg["journey_file"])
        logs = journey_data.get("logs", [])

        doc.add_heading(cfg["journey_name"], 1)
        doc.add_paragraph()

        master_json[f"journey_{journey_key}"] = {}

        for role, pid in [("REORDERER", cfg["reorderer"]), ("LAPSER", cfg["lapser"])]:
            log = get_log(logs, pid)
            if not log:
                add_para(doc, f"[No log found for {pid}]", color=CLR["lapse"])
                continue

            profile = get_persona_profile(cohort_lookup, pid)
            persona_rec = build_persona_json(
                pid, log, profile, cfg["decision_ticks"], cfg["product"]
            )

            role_label = f"{role}: {log.get('display_name')} ({pid})"
            master_json[f"journey_{journey_key}"][role.lower()] = persona_rec

            # ── Section heading ───────────────────────────────────────────────
            doc.add_heading(role_label, 2)
            add_para(
                doc,
                f"{'✅ Reordered' if log.get('reordered') else '❌ Did not reorder'}  ·  "
                f"Final decision: {(log.get('final_decision') or {}).get('decision', '—')}  ·  "
                f"Confidence: {(log.get('final_decision') or {}).get('confidence', '—')}",
                bold=True, size=10,
                color=CLR["reorder"] if log.get("reordered") else CLR["lapse"],
            )
            doc.add_paragraph()

            # ── Persona profile ───────────────────────────────────────────────
            if profile:
                doc.add_heading("Persona Profile", 3)
                demo = profile.get("demographic_anchor", {})
                derived = profile.get("derived_insights", {})
                beh = profile.get("behavioural_tendencies", {})
                attrs = profile.get("attributes", {})

                demo_rows = [
                    ("Name",              demo.get("name")),
                    ("Age",               demo.get("age")),
                    ("Gender",            demo.get("gender")),
                    ("City",              demo.get("location", {}).get("city")),
                    ("Urban tier",        demo.get("location", {}).get("urban_tier")),
                    ("Household",         f"{demo.get('household', {}).get('structure')} · size {demo.get('household', {}).get('size')}"),
                    ("Income bracket",    demo.get("household", {}).get("income_bracket")),
                    ("Education",         demo.get("education")),
                    ("Employment",        demo.get("employment")),
                    ("Decision style",    derived.get("decision_style")),
                    ("Trust anchor",      derived.get("trust_anchor")),
                    ("Risk appetite",     derived.get("risk_appetite")),
                    ("Primary value",     derived.get("primary_value_orientation")),
                    ("Coping mechanism",  derived.get("coping_mechanism")),
                    ("Price sensitivity", beh.get("price_sensitivity")),
                    ("Switching prop.",   beh.get("switching_propensity")),
                ]
                add_kv_table(doc, [(k, v) for k, v in demo_rows if v is not None], "Demographics & Decision Architecture")
                doc.add_paragraph()

                # Numeric attributes as compact table
                all_attrs = {}
                for group_name, group in attrs.items():
                    for attr_name, attr_val in group.items():
                        if "value" in attr_val:
                            all_attrs[f"{group_name}.{attr_name}"] = round(attr_val["value"], 4)
                if all_attrs:
                    add_para(doc, "Psychological attributes (0–1 scale):", bold=True, size=9, color=CLR["dim"])
                    attr_rows = sorted(all_attrs.items())
                    add_kv_table(doc, attr_rows)
                    doc.add_paragraph()

            # ── Decision snapshots ────────────────────────────────────────────
            doc.add_heading("Decision Snapshots", 3)
            snap_keys = list(persona_rec["decision_snapshots"].keys())
            labels = {
                snap_keys[0]: "First decision",
                snap_keys[-1]: "Final decision / reorder window",
            }
            for snap_key, snap_data in persona_rec["decision_snapshots"].items():
                tick = snap_data["tick"]
                add_decision_section(doc, snap_data, tick, labels.get(snap_key, f"tick {tick}"))

            # ── Deltas ────────────────────────────────────────────────────────
            doc.add_heading("What Changed Between Decisions", 3)
            add_delta_callout(doc, persona_rec["deltas"])

            # ── Post-decision stimuli ─────────────────────────────────────────
            doc.add_heading("Post-First-Decision Stimuli", 3)
            add_para(doc, "All stimuli fired after the first purchase decision, showing trust trajectory drivers:", size=9, color=CLR["dim"])
            add_stimulus_table(doc, persona_rec["post_first_decision_stimuli"])
            doc.add_paragraph()

            # ── Full JSON dump ────────────────────────────────────────────────
            doc.add_heading("Full Data Record (JSON)", 3)
            add_full_json_dump(doc, persona_rec)

            doc.add_page_break()

    # ── Save DOCX ─────────────────────────────────────────────────────────────
    doc.save(DOCX_OUT)
    print(f"✅  DOCX written to {DOCX_OUT}")

    # ── Save JSON ─────────────────────────────────────────────────────────────
    with open(JSON_OUT, "w", encoding="utf-8") as f:
        json.dump(master_json, f, indent=2, ensure_ascii=False)
    print(f"✅  JSON written to {JSON_OUT}")


if __name__ == "__main__":
    build()
